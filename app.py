import gradio as gr
import os
import json
import tempfile
from typing import List

try:
    from processors.doc_parser import extract_text_from_docx
    from processors.compliance_checker import check_document_completeness
    from processors.redflag_detector import detect_red_flags
    from processors.comment_inserter import insert_comments_in_docx
    from processors.clause_suggester import ClauseSuggester
except Exception as e:
    print("Import error in processors:", e)
    from docx import Document
    from io import BytesIO

    def _read_docx_bytes(file_obj):
        try:
            file_obj.file.seek(0)
            data = file_obj.file.read()
            return data
        except Exception:
            if isinstance(file_obj, str) and os.path.exists(file_obj):
                with open(file_obj, "rb") as fh:
                    return fh.read()
            raise

    def extract_text_from_docx(file_obj):
        data = _read_docx_bytes(file_obj)
        doc = Document(BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])

    def check_document_completeness(file_obj, process="Company Incorporation"):
        filename = getattr(file_obj, "name", str(file_obj))
        required = [
            "Register of Members and Directors",
            "Memorandum of Association",
            "UBO Declaration",
            "Articles of Association",
            "Board Resolution"
        ]
        matched = []
        fn = filename.lower()
        for doc in required:
            key = doc.split()[0].lower()
            if key in fn:
                matched.append(doc)
        missing = [d for d in required if d not in matched]
        return {"required_documents": len(required), "missing_documents": missing, "matched": matched}

    def detect_red_flags(text):
        issues = []
        if "uae" in text.lower() and "adgm" not in text.lower():
            issues.append({
                "section": "Jurisdiction Clause",
                "issue": "incorrect_jurisdiction",
                "severity": "High",
                "suggestion": "Update jurisdiction to ADGM Courts or include ADGM-specific clause."
            })
        if "sign" not in text.lower() and "signature" not in text.lower():
            issues.append({
                "section": "Signatory",
                "issue": "missing_signatory_section",
                "severity": "Medium",
                "suggestion": "Add signatory block with name, title, signature date and capacity."
            })
        return issues

    def insert_comments_in_docx(file_obj, issues):
        data = _read_docx_bytes(file_obj)
        temp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_in.write(data)
        temp_in.close()

        doc = Document(temp_in.name)
        if issues:
            doc.add_page_break()
            doc.add_paragraph("---- REVIEWER COMMENTS ----")
            for i, issue in enumerate(issues, start=1):
                suggestion_text = issue.get('suggestion', '')
                doc.add_paragraph(f"{i}. {issue['section']}: {issue['issue']} (Suggestion: {suggestion_text})")
        out_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(out_temp.name)
        return out_temp.name

    class ClauseSuggester:
        def __init__(self):
            self.suggestion_map = {
                "missing_signatory_section": "Add signatory block with name, title, signature date and capacity.",
                "incorrect_jurisdiction": "Update jurisdiction to ADGM Courts.",
                "ambiguous_language": "Clarify language to make obligations binding and explicit.",
                "missing_clause": "Include the required clause as per ADGM regulations.",
            }

        def suggest_clause_fix(self, issue_type: str) -> str:
            return self.suggestion_map.get(issue_type, "No suggestion available for this issue.")

def generate_json_report(process_name: str, filenames: List[str], checklist_result: dict, issues_by_file: List[dict]):
    total_uploaded = len(filenames)
    return {
        "process": process_name,
        "documents_uploaded": total_uploaded,
        "required_documents": checklist_result.get("required_documents", 0),
        "missing_documents": checklist_result.get("missing_documents", []),
        "matched_documents": checklist_result.get("matched", []),
        "issues_found": issues_by_file
    }

def process_documents(files):
    if not files:
        return {"error": "No files uploaded"}, []

    all_issues = []
    reviewed_paths = []
    saved_filenames = []

    suggester = ClauseSuggester()  # Instantiate clause suggester

    for f in files:
        try:
            text = extract_text_from_docx(f)
        except Exception as e:
            return {"error": f"Failed to parse uploaded file {getattr(f,'name',str(f))}: {e}"}, []

        checklist_result = check_document_completeness(f)
        issues = detect_red_flags(text)

        for issue in issues:
            # Use the issue key for suggestion mapping (which is now the 'issue' field)
            suggested_fix = suggester.suggest_clause_fix(issue.get("issue", ""))
            issue["suggestion"] = suggested_fix

        reviewed_path = insert_comments_in_docx(f, issues)
        reviewed_paths.append(reviewed_path)

        saved_filenames.append(getattr(f, "name", os.path.basename(reviewed_path)))

        for issue in issues:
            all_issues.append({
                "document": getattr(f, "name", os.path.basename(reviewed_path)),
                "section": issue.get("section"),
                "issue": issue.get("issue"),
                "severity": issue.get("severity"),
                "suggestion": issue.get("suggestion")
            })

    process_name = "Company Incorporation"

    summary = generate_json_report(process_name, saved_filenames, checklist_result, all_issues)

    try:
        os.makedirs("data", exist_ok=True)
        with open(os.path.join("data", "example_output.json"), "w", encoding="utf-8") as fh:
            json.dump(summary, fh, indent=2)
    except Exception:
        pass

    return summary, reviewed_paths


title = "ADGM Corporate Agent — Document Intelligence"
description = """
Upload one or more .docx documents (company formation documents, resolutions, declarations, etc.)
The app will analyze them for ADGM checklist completeness and red flags, generate a reviewed `.docx`
with appended reviewer comments and produce a structured JSON summary.
"""

with gr.Blocks() as demo:
    gr.Markdown(f"# {title}")
    gr.Markdown(description)

    with gr.Row():
        upload = gr.File(label="Upload .docx files", file_types=[".docx"], file_count="multiple")
    output_json = gr.JSON(label="Analysis Summary")
    output_files = gr.File(label="Reviewed Documents", file_count="multiple")

    btn = gr.Button("Submit")
    btn.click(fn=process_documents, inputs=upload, outputs=[output_json, output_files])

if __name__ == "__main__":
    demo.launch()
