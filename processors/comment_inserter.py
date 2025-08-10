from docx import Document

def add_inline_comments(docx_path, issues_found, output_path):
    """
    Opens a .docx file, inserts comments for each issue at the relevant paragraphs,
    and saves the reviewed document.
    """

    doc = Document(docx_path)

    # Simple example: Insert comment at the end of document for each issue.
    # (In real case, map issue locations more precisely)

    for issue in issues_found:
        comment_text = f"Issue: {issue['issue']}\nSuggestion: {issue.get('suggestion', 'No suggestion available')}"
        
        # For demonstration, add a new paragraph with comment text
        p = doc.add_paragraph()
        run = p.add_run(comment_text)
        font = run.font
        font.italic = True
        font.color.rgb = (255, 0, 0)  # red color (if supported)

        # For advanced inline comments, use Word comments feature if needed
        # e.g., using python-docx comments API or a library that supports it

    doc.save(output_path)
